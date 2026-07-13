from __future__ import annotations

import json
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path

from app.schemas import ExportScope, OutputFormat

MAX_FILENAME_PART_LENGTH = 120


@dataclass(frozen=True)
class TranscriptSegment:
    start: float
    end: float
    text: str


def format_timestamp(seconds: float) -> str:
    seconds = max(0, int(seconds))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def timed_line(segment: TranscriptSegment) -> str:
    return f"[{format_timestamp(segment.start)} -> {format_timestamp(segment.end)}] {segment.text}"


def segment_text(segments: list[TranscriptSegment], include_timestamps: bool) -> str:
    return "\n\n".join(timed_line(segment) if include_timestamps else segment.text for segment in segments).strip()


def srt_timestamp(seconds: float) -> str:
    milliseconds = max(0, int(round(seconds * 1000)))
    h = milliseconds // 3_600_000
    milliseconds %= 3_600_000
    m = milliseconds // 60_000
    milliseconds %= 60_000
    s = milliseconds // 1000
    ms = milliseconds % 1000
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def scope_sections(
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    scope: ExportScope,
) -> list[tuple[str, list[TranscriptSegment]]]:
    sections: list[tuple[str, list[TranscriptSegment]]] = []
    if scope in {ExportScope.raw, ExportScope.both}:
        sections.append(("原始转录文本", raw_segments))
    if scope in {ExportScope.polished, ExportScope.both} and polished_segments:
        sections.append(("整理后转录文本", polished_segments))
    if not sections:
        sections.append(("原始转录文本", raw_segments))
    return sections


def export_txt(
    path: Path,
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    include_timestamps: bool,
    scope: ExportScope,
) -> None:
    parts: list[str] = []
    for title, segments in scope_sections(raw_segments, polished_segments, scope):
        parts.append(f"{title}\n")
        parts.append("=" * len(title) + "\n\n")
        parts.append(segment_text(segments, include_timestamps) + "\n\n")
    path.write_text("".join(parts), encoding="utf-8")


def export_md(
    path: Path,
    title: str,
    metadata: dict[str, object],
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    include_timestamps: bool,
    scope: ExportScope,
) -> None:
    parts = [f"# {title}\n\n", "## 元数据\n\n"]
    for key, value in metadata.items():
        if value is not None:
            parts.append(f"- **{key}**: {value}\n")
    parts.append("\n")
    for section_title, segments in scope_sections(raw_segments, polished_segments, scope):
        parts.append(f"## {section_title}\n\n")
        if include_timestamps:
            parts.extend(
                f"- **{format_timestamp(segment.start)} -> {format_timestamp(segment.end)}** {segment.text}\n"
                for segment in segments
            )
            parts.append("\n")
        else:
            parts.append(segment_text(segments, include_timestamps=False) + "\n\n")
    path.write_text("".join(parts), encoding="utf-8")


def export_json(
    path: Path,
    metadata: dict[str, object],
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    include_timestamps: bool,
    scope: ExportScope,
) -> None:
    qwen_final_json = metadata.get("finalJson")
    if metadata.get("engine") == "qwen-audio" and isinstance(qwen_final_json, dict):
        with path.open("w", encoding="utf-8") as handle:
            json.dump(qwen_final_json, handle, ensure_ascii=False, indent=2)
        return

    payload = {
        "metadata": metadata,
        "segments": [asdict(segment) for segment in raw_segments],
        "rawText": segment_text(raw_segments, include_timestamps=False),
        "polishedText": segment_text(polished_segments or [], include_timestamps=False) if polished_segments else None,
        "parameters": {
            "includeTimestamps": include_timestamps,
            "exportScope": scope.value,
        },
    }
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)


def export_srt(path: Path, segments: list[TranscriptSegment]) -> None:
    parts: list[str] = []
    for index, segment in enumerate(segments, start=1):
        parts.append(f"{index}\n")
        parts.append(f"{srt_timestamp(segment.start)} --> {srt_timestamp(segment.end)}\n")
        parts.append(f"{segment.text}\n\n")
    path.write_text("".join(parts), encoding="utf-8")


def export_docx(
    path: Path,
    title: str,
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    include_timestamps: bool,
    scope: ExportScope,
) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, 0)
    for section_title, segments in scope_sections(raw_segments, polished_segments, scope):
        doc.add_heading(section_title, 1)
        for segment in segments:
            paragraph = doc.add_paragraph()
            if include_timestamps:
                run = paragraph.add_run(f"[{format_timestamp(segment.start)} -> {format_timestamp(segment.end)}] ")
                run.bold = True
            paragraph.add_run(segment.text)
    doc.save(path)


def export_transcript(
    output_dir: Path,
    base_name: str,
    formats: list[OutputFormat],
    raw_segments: list[TranscriptSegment],
    polished_segments: list[TranscriptSegment] | None,
    include_timestamps: bool,
    scope: ExportScope,
    metadata: dict[str, object] | None = None,
) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    title = f"{base_name} 转录结果"
    export_metadata = {
        "title": title,
        "createdAt": datetime.now().isoformat(timespec="seconds"),
        **(metadata or {}),
    }
    date_part = datetime.now().strftime("%Y%m%d")
    language_part = safe_filename_part(str(export_metadata.get("language") or "auto"))
    model_part = safe_filename_part(str(export_metadata.get("transcriptionModel") or "model"))
    writers = {
        OutputFormat.txt: lambda path: export_txt(path, raw_segments, polished_segments, include_timestamps, scope),
        OutputFormat.md: lambda path: export_md(
            path,
            title,
            export_metadata,
            raw_segments,
            polished_segments,
            include_timestamps,
            scope,
        ),
        OutputFormat.json: lambda path: export_json(
            path,
            export_metadata,
            raw_segments,
            polished_segments,
            include_timestamps,
            scope,
        ),
        OutputFormat.srt: lambda path: export_srt(
            path,
            polished_segments if scope == ExportScope.polished and polished_segments else raw_segments,
        ),
        OutputFormat.docx: lambda path: export_docx(path, title, raw_segments, polished_segments, include_timestamps, scope),
    }

    paths: list[Path] = []
    timestamp_suffix = "timed" if include_timestamps else "plain"
    scope_suffix = scope.value
    for fmt in formats:
        if fmt == OutputFormat.srt and not include_timestamps:
            continue
        path = output_dir / f"{safe_filename_part(base_name)}_{date_part}_{language_part}_{model_part}_{scope_suffix}_{timestamp_suffix}.{fmt.value}"
        writers[fmt](path)
        paths.append(path)
    return paths


def safe_filename_part(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in value)
    filename_part = cleaned.strip("_") or "item"
    return filename_part[:MAX_FILENAME_PART_LENGTH].rstrip("_") or "item"
